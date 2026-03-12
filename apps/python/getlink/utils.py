import psycopg2
from psycopg2 import extras
import pandas as pd
from docx import Document
from fpdf import FPDF
import io

def get_db_connection(host, port, user, password, dbname):
    try:
        conn = psycopg2.connect(
            host=host,
            port=port,
            user=user,
            password=password,
            dbname=dbname,
            connect_timeout=5
        )
        return conn, None
    except Exception as e:
        return None, str(e)

def fetch_user_links(conn, levels, base_url):
    try:
        cursor = conn.cursor(cursor_factory=extras.RealDictCursor)
        
        # Prepare levels for SQL IN clause
        levels_str = ",".join([str(lvl) for lvl in levels])
        
        query = f"""
            SELECT u.nama, u.username, k.nama as kelas_nama, k.tingkat
            FROM users u
            JOIN kelas k ON u.kelas_id = k.id
            WHERE k.tingkat IN ({levels_str}) AND u.deleted_at IS NULL
            ORDER BY k.tingkat ASC, k.nama ASC, u.nama ASC
        """
        
        cursor.execute(query)
        rows = cursor.fetchall()
        
        data = []
        for row in rows:
            username = row['username']
            # Ensure base_url doesn't end with / if username starts with /
            clean_base_url = base_url.rstrip('/')
            link = f"{clean_base_url}/{username}"
            data.append({
                'Nama': row['nama'],
                'Username': username,
                'Kelas': row['kelas_nama'],
                'Tingkat': row['tingkat'],
                'Link': link
            })
            
        return data, None
    except Exception as e:
        return None, str(e)

def export_to_txt(data):
    output = io.StringIO()
    for item in data:
        output.write(f"{item['Nama']} - {item['Link']}\n")
    return output.getvalue().encode('utf-8')

def export_to_docx(data):
    doc = Document()
    doc.add_heading('Daftar Link Profil Grafikarsa', 0)
    
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Nama'
    hdr_cells[1].text = 'Link Profil'
    
    for item in data:
        row_cells = table.add_row().cells
        row_cells[0].text = item['Nama']
        row_cells[1].text = item['Link']
        
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()

def export_to_pdf(data):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", 'B', 16)
    pdf.cell(200, 10, txt="Daftar Link Profil Grafikarsa", ln=True, align='C')
    pdf.ln(10)
    
    pdf.set_font("Arial", size=10)
    
    # Table Header
    pdf.set_fill_color(240, 240, 240)
    pdf.cell(80, 10, "Nama", 1, 0, 'C', 1)
    pdf.cell(110, 10, "Link Profil", 1, 1, 'C', 1)
    
    for item in data:
        # Handle long names by shortening or using multi_cell if needed
        # For simplicity, we just use cell here
        name = item['Nama'][:40]
        link = item['Link']
        pdf.cell(80, 8, name, 1)
        pdf.cell(110, 8, link, 1, 1)
        
    return pdf.output(dest='S').encode('latin-1')

def export_to_csv(data):
    df = pd.DataFrame(data)
    # Only export Nama and Link as requested implicitly (or all)
    # The requirement says "Nama - Link" in TXT, so let's stick to that for CSV too if possible
    # but CSV is usually better with more info.
    return df[['Nama', 'Link']].to_csv(index=False).encode('utf-8')

def process_batch_file(uploaded_file, domain="https://grafikarsa.com"):
    try:
        # Determine file type
        file_name = uploaded_file.name.lower()
        
        if file_name.endswith('.csv'):
            # Read CSV
            content = uploaded_file.read()
            uploaded_file.seek(0)
            df = pd.read_csv(io.BytesIO(content))
            # Heuristic for no-header CSV
            nis_col = None
            for col in df.columns:
                if 'nis' in str(col).lower():
                    nis_col = col
                    break
            if nis_col is None:
                df_no_hdr = pd.read_csv(io.BytesIO(content), header=None)
                for col in df_no_hdr.columns:
                    if df_no_hdr[col].astype(str).str.match(r'^\d+$').sum() > len(df_no_hdr) / 2:
                        nis_col = col
                        df = df_no_hdr
                        break
        elif file_name.endswith(('.xlsx', '.xls')):
            # Read Excel
            df = pd.read_excel(uploaded_file)
            nis_col = None
            for col in df.columns:
                if 'nis' in str(col).lower():
                    nis_col = col
                    break
        else:
            return None, "Format file tidak didukung. Gunakan CSV atau Excel (.xlsx, .xls)."
        
        if nis_col is not None:
            # Clean domain
            clean_domain = domain.rstrip('/')
            
            # Add Link Profil column
            # Ensure NIS is treated as string and cleaned
            df['Link Profil'] = df[nis_col].apply(lambda x: f"{clean_domain}/{str(x).strip().split('.')[0] if '.' in str(x) else str(x).strip()}")
            return df, None
        else:
            return None, "Kolom NIS tidak ditemukan (Coba pastikan ada kolom bernama 'nis' atau data NIS berupa angka)."
            
    except Exception as e:
        return None, str(e)
